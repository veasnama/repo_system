
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def parse_table_data(text):
    # Split text into lines and process
    lines = [line.strip() for line in text.strip().split('\n') if line.strip()]
    tables_data = []
    current_rows = []

    for i, line in enumerate(lines):
        # Start of a new table
        if line.startswith('Temperature') and 'Voltage' in line and 'Tx Bias' in line:
            # Skip the units line and dashed line
            if i + 2 < len(lines) and lines[i + 1].startswith(('(C)', '--------------')):
                i += 2
            if current_rows:
                tables_data.append((['Label', 'Temperature', 'Voltage', 'Tx Bias', 'Tx Power', 'Rx Power'], current_rows))
            current_rows = []
            continue
        
        # Collect data rows
        if (line.startswith(('Value', 'Status', 'High Alarm', 'High Warning', 'Low Warning', 'Low Alarm')) or not line):
            if not line:
                if current_rows:
                    tables_data.append((['Label', 'Temperature', 'Voltage', 'Tx Bias', 'Tx Power', 'Rx Power'], current_rows))
                    current_rows = []
                continue
            columns = re.split(r'\s+', line.strip())
            # Handle multi-word labels
            if columns[0] in ('High', 'Low'):
                label = f"{columns[0]} {columns[1]}"
                data_columns = columns[2:]  # Start data after the two-word label
            else:
                label = columns[0]
                data_columns = columns[1:]  # Start data after the single-word label
            
            if len(data_columns) >= 5:  # Ensure enough columns for sensor data
                filtered_columns = [label] + data_columns[:5]  # Label + Temperature to Rx Power
                current_rows.append(filtered_columns)
    
    # Append the last table if exists
    if current_rows:
        tables_data.append((['Label', 'Temperature', 'Voltage', 'Tx Bias', 'Tx Power', 'Rx Power'], current_rows))
    
    print(f"Parsed tables: {len(tables_data)}")
    for headers, rows in tables_data:
        print(f"Headers: {headers}")
        for row in rows:
            print(f"Row: {row}")
    return tables_data

def create_word_table(headers_rows_list, output_file):
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('System Resource Tables', level=1)
    
    # Create a table for each sensor data block
    for headers, rows in headers_rows_list:
        if not rows:
            print("Warning: No data rows for this table")
            continue
        
        # Create table with headers + rows
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set header row
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(12)
        
        # Populate data rows
        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_data if cell_data else ''
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.font.size = Pt(10)
        
        # Adjust column widths
        for j in range(len(headers)):
            column = table.columns[j]
            max_width = max((len(str(cell.text)) for cell in column.cells), default=1)
            for cell in column.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.size = Pt(10)
            column.width = Inches(max(0.5, min(1.0, max_width * 0.05)))
    
    # Save the document
    doc.save(output_file)
    print(f"Word document saved as '{output_file}'")

# Example text data
text_data = """
SunOS DR-CBSCMS-PRM02 5.11 11.4.77.182.2 sun4v sparc sun4v non-virtualized

Installation directory: /usr/lib/ssm/fwupdate/qlogic
Working dir: /var/explorer/output/explorer.86e22332.DR-CBSCMS-PRM02-2025.04.01.02.54/rda/work
--------------------------------------------------------------------------------
HBA/OS Instance 0/0: QLE2742 Port 1 WWPN 21-00-00-24-FF-19-84-C6 PortID 01-05-00
Link: Online
--------------------------------------------------------------------------------
--------------------------------------------------------------------------------
Media Information
--------------------------------------------------------------------------------
              Vendor: FINISAR CORP.
                Type: 800-M5-SN-S
         Part Number: FTLF8532P4BCV-QL
               Speed: 3200 MBytes/Sec 1600 MBytes/Sec 800 MBytes/Sec
            Revision: A
       Serial Number: PY80WQP
QLogic SFP Installed: Yes
--------------------------------------------------------------------------------
                 Temperature   Voltage     Tx Bias    Tx Power    Rx Power
                     (C)         (V)        (mA)        (mW)        (mW)
              --------------   -------     -------    --------    --------
        Value      46.09         3.45        6.56      0.6409      0.6417
       Status     Normal       Normal      Normal      Normal      Normal
   High Alarm      75.00         3.60       12.00      1.9953      1.9953
 High Warning      70.00         3.50       11.50      1.5849      1.5849
  Low Warning       0.00         3.10        2.00      0.1585      0.0158
    Low Alarm      -5.00         3.00        1.00      0.1259      0.0100
--------------------------------------------------------------------------------
HBA/OS Instance 1/1: QLE2742 Port 2 WWPN 21-00-00-24-FF-19-84-C7 PortID 02-05-00
Link: Online
--------------------------------------------------------------------------------
--------------------------------------------------------------------------------
Media Information
--------------------------------------------------------------------------------
              Vendor: FINISAR CORP.
                Type: 800-M5-SN-S
         Part Number: FTLF8532P4BCV-QL
               Speed: 3200 MBytes/Sec 1600 MBytes/Sec 800 MBytes/Sec
            Revision: A
       Serial Number: PY80X3T
QLogic SFP Installed: Yes
--------------------------------------------------------------------------------
                 Temperature   Voltage     Tx Bias    Tx Power    Rx Power
                     (C)         (V)        (mA)        (mW)        (mW)
              --------------   -------     -------    --------    --------
        Value      47.19         3.42        6.56      0.3315      0.7046
       Status     Normal       Normal      Normal      Normal      Normal
   High Alarm      75.00         3.60       12.00      1.9953      1.9953
 High Warning      70.00         3.50       11.50      1.5849      1.5849
  Low Warning       0.00         3.10        2.00      0.1585      0.0158
    Low Alarm      -5.00         3.00        1.00      0.1259      0.0100
--------------------------------------------------------------------------------
HBA/OS Instance 2/2: QLE2742 Port 1 WWPN 21-00-00-24-FF-19-86-D0 PortID 01-06-00
Link: Online
--------------------------------------------------------------------------------
--------------------------------------------------------------------------------
Media Information
--------------------------------------------------------------------------------
              Vendor: FINISAR CORP.
                Type: 800-M5-SN-S
         Part Number: FTLF8532P4BCV-QL
               Speed: 3200 MBytes/Sec 1600 MBytes/Sec 800 MBytes/Sec
            Revision: A
       Serial Number: PY80WLL
QLogic SFP Installed: Yes
--------------------------------------------------------------------------------
                 Temperature   Voltage     Tx Bias    Tx Power    Rx Power
                     (C)         (V)        (mA)        (mW)        (mW)
              --------------   -------     -------    --------    --------
        Value      47.12         3.41        6.56      0.6875      0.5973
       Status     Normal       Normal      Normal      Normal      Normal
   High Alarm      75.00         3.60       12.00      1.9953      1.9953
 High Warning      70.00         3.50       11.50      1.5849      1.5849
  Low Warning       0.00         3.10        2.00      0.1585      0.0158
    Low Alarm      -5.00         3.00        1.00      0.1259      0.0100
--------------------------------------------------------------------------------
HBA/OS Instance 3/3: QLE2742 Port 2 WWPN 21-00-00-24-FF-19-86-D1 PortID 02-06-00
Link: Online
--------------------------------------------------------------------------------
--------------------------------------------------------------------------------
Media Information
--------------------------------------------------------------------------------
              Vendor: FINISAR CORP.
                Type: 800-M5-SN-S
         Part Number: FTLF8532P4BCV-QL
               Speed: 3200 MBytes/Sec 1600 MBytes/Sec 800 MBytes/Sec
            Revision: A
       Serial Number: PY80X30
QLogic SFP Installed: Yes
--------------------------------------------------------------------------------
                 Temperature   Voltage     Tx Bias    Tx Power    Rx Power
                     (C)         (V)        (mA)        (mW)        (mW)
              --------------   -------     -------    --------    --------
        Value      42.51         3.42        6.54      0.6686      0.6359
       Status     Normal       Normal      Normal      Normal      Normal
   High Alarm      75.00         3.60       12.00      1.9953      1.9953
 High Warning      70.00         3.50       11.50      1.5849      1.5849
  Low Warning       0.00         3.10        2.00      0.1585      0.0158
    Low Alarm      -5.00         3.00        1.00      0.1259      0.0100
"""

# Parse and create tables
tables_data = parse_table_data(text_data)
create_word_table(tables_data, "ParseQLogic.docx")
