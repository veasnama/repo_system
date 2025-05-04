from docx import Document
from docx.shared import Pt, Inches
import re

def parse_table_data(text):
    # Split text into lines and filter out empty lines
    lines = [line.strip() for line in text.strip().split('\n') if line.strip()]
    
    # Extract headers (first line) and exclude "FLAGS" and "CONS"
    headers = lines[0].split()
    exclude_cols = ["FLAGS", "CONS"]
    filtered_headers = [h for h in headers if h not in exclude_cols]
    print(f"Filtered headers: {filtered_headers}")  # Debug: Confirm headers
    
    # Process rows, handling variable spacing and missing columns
    rows = []
    for line in lines[1:]:
        columns = re.split(r'\s+', line.strip())
        # Pad with empty strings if fewer columns than original headers
        if len(columns) < len(headers):
            columns.extend([''] * (len(headers) - len(columns)))
        # Truncate if more columns than headers
        if len(columns) > len(headers):
            columns = columns[:len(headers)]
        # Filter out columns corresponding to "FLAGS" and "CONS"
        filtered_columns = [col for i, col in enumerate(columns) if headers[i] not in exclude_cols]
        rows.append(filtered_columns)
    
    print(f"Parsed rows: {len(rows)}")  # Debug: Confirm number of rows
    for row in rows:
        print(f"Row: {row}")  # Debug: Inspect each row
    return filtered_headers, rows

def create_word_table(headers, rows, output_file):
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('System Resource Table', level=1)
    
    # Check if rows are empty
    if not rows:
        print("Warning: No data rows to add to the table")
    
    # Create a table with headers + rows
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'  # Apply a grid style
    table.autofit = True
    
    # Set header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        # Format header: bold, centered, larger font
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1  # Center
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(12)
    
    # Populate data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_data if cell_data else ''  # Ensure empty cells are handled
            # Format cell: centered, regular font
            paragraph = cell.paragraphs[0]
            paragraph.alignment = 1  # Center
            run = paragraph.runs[0]
            run.font.size = Pt(10)
    
    # Adjust column widths (optional, set to fit content)
    for column in table.columns:
        column.width = Inches(1.0)  # Adjust as needed
    
    # Save the document
    doc.save(output_file)
    print(f"Word document saved as '{output_file}'")

# Example text data
text_data = """
NAME             STATE      FLAGS   CONS    VCPU  MEMORY   UTIL  NORM  UPTIME
primary          active     -n-cv-  UART    16    32G      7.0%  6.9%  45d 8h 6m
DR-PRD02-CBSAPP01 active     -n----  5004    16    64G      0.3%  0.3%  45d 7h 55m
DR-PRD02-CBSDB01 active     -n----  5014    64    164G     3.9%  3.9%  45d 8h 3m
DR-PRD02-CBSDB02 bound      ------  5006    16    100G
DR-PRD02-CBSGW01 active     -n----  5012    8     16G      0.6%  0.6%  45d 8h 6m
DR-PRD02-CBSRPT01 bound      ------  5013    16    32G
DR-PRD02-CIFTPDB01 active     -n----  5003    8     64G      5.0%  5.0%  45d 7h 54m
DR-PRD02-CMSAPP01 active     -n----  5002    16    32G      0.4%  0.4%  45d 8h 2m
DR-PRD02-CMSAPP02 active     -n----  5009    16    32G      0.4%  0.4%  45d 8h 2m
DR-PRD02-CMSAPP03 active     -n----  5010    16    32G      0.3%  0.3%  45d 8h 2m
DR-PRD02-CMSDB01 active     -n----  5000    16    100G     5.0%  5.0%  45d 8h 3m
DR-PRD02-CMSDB02 active     -n----  5001    16    48G      5.2%  5.1%  45d 8h 3m
"""

# DR-PRD02-OBPMDB01 active     -n----  5005    16    64G      2.5%  2.5%  45d 8h 3m
# Parse and create table
