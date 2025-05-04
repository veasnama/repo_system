from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def parse_table_data(text):
    # Split text into lines and filter out empty lines
    lines = [line.strip() for line in text.strip().split('\n') if line.strip()]
    
    # Extract headers (first line) and exclude "DEDUP" and "ALTROOT"
    headers = lines[0].split()
    exclude_cols = ["DEDUP", "ALTROOT"]
    filtered_headers = [h for h in headers if h not in exclude_cols]
    print(f"Filtered headers: {filtered_headers}")  # Debug: Confirm headers
    
    # Process rows, handling variable spacing and missing columns
    rows = []
    for line in lines[1:]:
        columns = re.split(r'\s+', line.strip())
        # Pad with empty strings if fewer columns than original headers
        if len(columns) < len(headers):
            columns.extend([''] * (len(headers) - len(columns)))
        # Truncate if more columns than headers
        if len(columns) > len(headers):
            columns = columns[:len(headers)]
        # Filter out columns corresponding to "DEDUP" and "ALTROOT"
        filtered_columns = [col for i, col in enumerate(columns) if headers[i] not in exclude_cols]
        rows.append(filtered_columns)
    
    print(f"Parsed rows: {len(rows)}")  # Debug: Confirm number of rows
    for row in rows:
        print(f"Row: {row}")  # Debug: Inspect each row
    return filtered_headers, rows

def create_word_table(headers, rows, output_file):
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('System Resource Table', level=1)
    
    # Check if rows are empty
    if not rows:
        print("Warning: No data rows to add to the table")
    
    # Create a table with headers + rows
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'  # Apply a grid style
    table.autofit = False  # Disable auto-fit to respect manual widths
    
    # Set header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        # Format header: bold, centered, larger font
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(12)
    
    # Populate data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_data if cell_data else ''  # Ensure empty cells are handled
            # Format cell: centered, regular font
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(10)
    
    # Adjust column widths
    for j in range(len(headers)):
        column = table.columns[j]
        if headers[j] == "NAME":
            # Set a larger fixed width for the "NAME" column
            column.width = Inches(2.5)  # 2.5 inches for "NAME"
        else:
            # Set a smaller width for other columns based on content
            max_width = max((len(str(cell.text)) for cell in table.columns[j].cells), default=1)
            for cell in column.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.size = Pt(10)  # Ensure consistent font size for width calculation
            column.width = Inches(max(0.5, min(1.0, max_width * 0.05)))  # Smaller range for other columns (0.5 to 1 inch)
    
    # Save the document
    doc.save(output_file)
    print(f"Word document saved as '{output_file}'")

# Example text data
text_data = """
NAME         SIZE  ALLOC   FREE  CAP  DEDUP  HEALTH  ALTROOT
rpool       1.09T   768G   349G  68%  1.01x  ONLINE  -
"""

# ldompool    2.17T  1.11T  1.06T  51%  1.00x  ONLINE  -
# ldompool02  1016G   193G   823G  18%  1.00x  ONLINE  -
# Parse and create table
headers, rows = parse_table_data(text_data)
create_word_table(headers, rows, "parse_zpool.docx")
