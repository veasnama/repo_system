from docx import Document
import re

def parse_log_data(log_text):
    # Initialize list to store parsed entries
    entries = []
    current_entry = {}
    
    # Split log into individual events using the separator line
    events = re.split(r'\n*-{10,} -{10,}  -{10,} -{5,}\n', log_text.strip())
    
    for event in events:
        if not event.strip():
            continue
            
        # Extract Time
        time_match = re.search(r'(\w{3} \d{2} \d{2}:\d{2}:\d{2})', event)
        if time_match:
            current_entry['Time'] = time_match.group(1)
        else:
            print(f"Warning: No Time found in event: {event[:100]}...")  # Debug output
            continue
            
        # Extract Problem class
        problem_class_match = re.search(r'Problem class\s*:\s*(.+)', event)
        if problem_class_match:
            current_entry['Problem class'] = problem_class_match.group(1)
        else:
            print(f"Warning: No Problem class found in event: {event[:100]}...")  # Debug output
            continue
            
        # Extract Server_Name
        server_name_match = re.search(r'Server_Name\s*:\s*(.+)', event)
        if server_name_match:
            current_entry['Server_Name'] = server_name_match.group(1)
        else:
            print(f"Warning: No Server_Name found in event: {event[:100]}...")  # Debug output
            continue
            
        # Extract Description
        description_match = re.search(r'Description\s*:\s*(.+?)(?:\nResponse|$)', event, re.DOTALL)
        if description_match:
            current_entry['Description'] = description_match.group(1).strip()
        else:
            print(f"Warning: No Description found in event: {event[:100]}...")  # Debug output
            continue
            
        # Extract Action (capture all lines until end of event or next section)
        action_match = re.search(r'Action\s*:\s*(.+?)(?=\n{2,}|$)', event, re.DOTALL)
        if action_match:
            # Clean up Action field: replace multiple newlines/spaces with single space
            action_text = re.sub(r'\s+', ' ', action_match.group(1)).strip()
            current_entry['Action'] = action_text
        else:
            print(f"Warning: No Action found in event: {event[:100]}...")  # Debug output
            continue
            
        # Only append if all required fields are present
        required_fields = ['Time', 'Problem class', 'Server_Name', 'Description', 'Action']
        if all(field in current_entry for field in required_fields):
            entries.append(current_entry.copy())
        else:
            print(f"Warning: Incomplete entry skipped: {current_entry}")
            
        current_entry.clear()
    
    return entries

def create_docx(entries, output_file):
    # Create new Document
    doc = Document()
    
    # Add title
    doc.add_heading('Disk Error Report', 0)
    
    # Add bullet points for each entry
    for entry in entries:
        doc.add_heading(f'Event at {entry["Time"]}', level=1)
        for key, value in entry.items():
            if key != 'Time':
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(value)
                p.style = 'List Bullet'
    
    # Save the document
    doc.save(output_file)

# Read log data from file content
log_data = """--------------- ------------------------------------  -------------- ---------
TIME            EVENT-ID                              MSG-ID         SEVERITY
--------------- ------------------------------------  -------------- ---------
Oct 30 06:45:40 cec62661-80cc-4b15-b656-b46edd64d7ac  DISK-8000-FA   Critical
Problem Status            : isolated
Diag Engine               : eft / 1.16
System
    Manufacturer          : Oracle Corporation
    Name                  : SPARC T8-1
    Part_Number           : 7113718-12922
    Serial_Number         : 2052NMC00A
System Component
    Firmware_Manufacturer : Oracle Corporation
    Firmware_Version      : (OBP)4.43.10,(HC)1.11.10,(HV)1.20.6.c
    Firmware_Release      : (OBP)2023.03.09,(HC)2023.03.09,(HV)2023.03.09
    Host_ID               : 86e0bd62
    Server_Name           : FCUBS-DC-PS02
    Virtualization_Type   : logical-domain
    Virtualization_Roles  : control-domain,io-domain,root-domain,service-domain
    Parent_Type           : non-virtualized
    OS_Version            : 11.4.31.88.5
----------------------------------------
Suspect 1 of 1 :
   Problem class : fault.io.scsi.disk.probe-failure
   Certainty   : 100%
   Affects     : dev:///:devid=id1,sd@n5000c500ca3467b7//scsi_vhci/disk@g5000c500ca3467b7
   Status      : faulted and taken out of service
   FRU
     Status           : faulty
     Location         : "/SYS/DBP/HDD4"
     Manufacturer     : SEAGATE
     Name             : ST1200IN9SUN1.2T
     Part_Number      : SEAGATE-ST1200IN9SUN1.2T
     Revision         : ORA9
     Serial_Number    : 002017L7N88C--------WFK7N88C
     Chassis
        Manufacturer  : Oracle Corporation
        Name          : SPARC T8-1
        Part_Number   : 7113718-12922
        Serial_Number : 2052NMC00A
Description : A device could not be accessed.
Response    : The device may be offlined or degraded. An attempt will be made
              to activate a hot spare if available.
Impact      : The device is not accessible. The service may have been lost or
              degraded.
Action      : Please refer to the associated reference document at
              http://support.oracle.com/msg/DISK-8000-FA for the latest service
              procedures and policies regarding this diagnosis.
--------------- ------------------------------------  -------------- ---------
Oct 30 06:45:09 e2e0e4bb-0fd1-44f1-9ec4-b97d7b09c434  DISK-8000-3E   Critical
Problem Status            : isolated
Diag Engine               : eft / 1.16
System
    Manufacturer          : Oracle Corporation
    Name                  : SPARC T8-1
    Part_Number           : 7113718-12922
    Serial_Number         : 2052NMC00A
System Component
    Firmware_Manufacturer : Oracle Corporation
    Firmware_Version      : (OBP)4.43.10,(HC)1.11.10,(HV)1.20.6.c
    Firmware_Release      : (OBP)2023.03.09,(HC)2023.03.09,(HV)2023.03.09
    Host_ID               : 86e0bd62
    Server_Name           : FCUBS-DC-PS02
    Virtualization_Type   : logical-domain
    Virtualization_Roles  : control-domain,io-domain,root-domain,service-domain
    Parent_Type           : non-virtualized
    OS_Version            : 11.4.31.88.5
----------------------------------------
Suspect 1 of 1 :
   Problem class : fault.io.scsi.cmd.disk.dev.rqs.derr
   Certainty   : 100%
   Affects     : dev:///:devid=id1,sd@n5000c500ca3467b7//scsi_vhci/disk@g5000c500ca3467b7
   Status      : faulted and taken out of service
   FRU
     Status           : faulty
     Location         : "/SYS/DBP/HDD4"
     Manufacturer     : SEAGATE
     Name             : ST1200IN9SUN1.2T
     Part_Number      : SEAGATE-ST1200IN9SUN1.2T
     Revision         : ORA9
     Serial_Number    : 002017L7N88C--------WFK7N88C
     Chassis
        Manufacturer  : Oracle Corporation
        Name          : SPARC T8-1
        Part_Number   : 7113718-12922
        Serial_Number : 2052NMC00A
Description : A non-recoverable hardware failure was detected by the device.
Response    : The device may be offlined or degraded. A hot-spare disk may have
              been activated.
Impact      : The device has failed. The service may have been lost or
              degraded.
Action      : Please refer to the associated reference document at
              http://support.oracle.com/msg/DISK-8000-3E for the latest service
              procedures and policies regarding this diagnosis.
"""

# Parse the log data
parsed_entries = parse_log_data(log_data)

# Debug: Print parsed entries to verify, especially Action field with URLs
print(f"Parsed {len(parsed_entries)} entries:")
for i, entry in enumerate(parsed_entries, 1):
    print(f"Entry {i}:")
    for key, value in entry.items():
        print(f"  {key}: {value}")
    # Check if Action contains expected URL
    if 'Action' in entry and 'http://support.oracle.com/msg/DISK-8000' in entry['Action']:
        print(f"  [Confirmed: Action contains expected URL]")
    else:
        print(f"  [Warning: Action may not contain expected URL]")

# Create DOCX file
create_docx(parsed_entries, 'disk_error_report.docx')
